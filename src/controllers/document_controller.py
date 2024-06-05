from flask import request, Response, json, Blueprint, render_template, make_response
from src.models.document_model import Document
from src.middlewares import authentication_required
from src import db
import pdfkit
from docx import Document as DocumentLib
from io import BytesIO
import platform

# Get the operating system name
os_name = platform.system()

# document controller blueprint to be registered with api blueprint
documents = Blueprint("documents", __name__)

# Get document by id
@documents.route('/<int:document_id>', methods=['GET'])
@authentication_required
def get_document(document_id):
    try:
        document_obj = Document.query.get(document_id)
        if document_obj.status == False:
            available_count = document_obj.available_sentences
            content = document_obj.content
            generated_essay = {}
            generated_essay["id"] = document_id
            generated_essay["title"] = content['title']
            generated_essay["sections"] = []
            generated_essay['reference_result'] = content['reference_result']
                
            count = 0
            for section in content["sections"]:
                essay_section = {
                    "section_title": section["section_title"],
                    "section_content": [],
                    "subsections": [],
                    'summarize_content': section['summarize_content']
                }                    
                for section_content in section["section_content"]:
                    essay_section["section_content"].append(section_content)    
                    count = count + 1
                    if count == available_count:
                        generated_essay["sections"].append(essay_section)
                        return Response(
                            response=json.dumps({
                                "status": True,
                                "data": {
                                    "essay": generated_essay,
                                }
                            }),
                            status=200,
                            mimetype="application/json"
                        )
                
                for subsection in section["subsections"]:
                    essay_subsection = {
                        "subsection_title": subsection["subsection_title"],
                        "subsection_content": [],
                        'summarize_content': subsection['summarize_content']
                    }
                    for subsection_content in subsection["subsection_content"]:
                        essay_subsection["subsection_content"].append(subsection_content)
                        count = count + 1
                        if count == available_count:
                            essay_section["subsections"].append(essay_subsection)
                            generated_essay["sections"].append(essay_section)
                            return Response(
                                response=json.dumps({
                                    "status": True,
                                    "data": {
                                        "essay": generated_essay,
                                    }
                                }),
                                status=200,
                                mimetype="application/json"
                            )
                    essay_section["subsections"].append(essay_subsection)
                generated_essay["sections"].append(essay_section)
                
            generated_essay["references"] = content["references"]
            
            return Response(
                response=json.dumps({
                    "status": True,
                    "data": {
                        "essay": generated_essay,
                    }
                }),
                status=200,
                mimetype="application/json"
            )

        else:
            return Response(
                    response=json.dumps({
                        "status": True,
                        "data": {
                            "essay": {
                                **document_obj.content,
                                "id": document_obj.id
                            }
                        }
                    }),
                    status=200,
                    mimetype="application/json"
                )
    except Exception as e:
        return Response(
                response=json.dumps({
                    "status": False,
                    "message": "Error Occured",
                    "error": str(e)
                }),
                status=500,
                mimetype="application/json"
            )
        
# Get documents by user_id
@documents.route('', methods=['GET'])
@authentication_required
def get_documents():
    try:
        user_id = request.user_id
        document_objs = Document.query.filter_by(user_id = user_id).order_by(Document.created_at.desc()).with_entities(Document.id, Document.title)
        documents = [{"id": doc.id, "title": doc.title} for doc in document_objs]
        
        return Response(
                response=json.dumps({
                    "status": True,
                    "data": {
                        "documents": documents
                    }
                }),
                status=200,
                mimetype="application/json"
            )
    except Exception as e:
        return Response(
                response=json.dumps({
                    "status": False,
                    "message": "Error Occured",
                    "error": str(e)
                }),
                status=500,
                mimetype="application/json"
            )
        
# Delete document by id
@documents.route('/<int:document_id>', methods=['DELETE'])
@authentication_required
def delete_document(document_id):
    try:
        document_obj = Document.query.get(document_id)
        if document_obj:
            db.session.delete(document_obj)
            db.session.commit()
            return Response(
                response=json.dumps({
                    "status": True,
                    "message": "Document has been deleted successfully"
                }),
                status=200,
                mimetype="application/json"
            )
        else:
            return Response(
                response=json.dumps({
                    "status": False,
                    "message": "Document not found"
                }),
                status=404,
                mimetype="application/json"
            )
    except Exception as e:
        return Response(
                response=json.dumps({
                    "status": False,
                    "message": "Error Occured",
                    "error": str(e)
                }),
                status=500,
                mimetype="application/json"
            )
        
@documents.route('/<int:document_id>/export', methods=['POST'])
@authentication_required
def export_document(document_id):
    try:
        data = request.json
        file_tyoe = data["type"]
        
        document_obj = Document.query.get(document_id)
        if document_obj:
            if file_tyoe == "PDF":
                html = render_template('essay.html', essay=document_obj.content)
                if os_name == "Windows":
                    config = pdfkit.configuration(wkhtmltopdf="C:\\Program Files\\wkhtmltopdf\\bin\\wkhtmltopdf.exe")
                else: config = pdfkit.configuration(wkhtmltopdf="/usr/bin/wkhtmltopdf")
                pdf = pdfkit.from_string(html, False, configuration=config)
                response = make_response(pdf)
                response.headers['Content-Type'] = 'application/pdf'
                response.headers['Content-Disposition'] = 'attachment; filename="{}.pdf"'.format(document_obj.title)
                return response
                
            else:
                document = DocumentLib()
                document.add_heading(f"{document_obj.title}", level=1).alignment = 1
                essay = document_obj.content
                for section in essay["sections"]:
                    document.add_heading(f"{section["section_title"]}", level=3)
                    con_text = ""
                    for con in section["section_content"]:
                        con_text = con_text + f"{con["sentence"]} "
                    if con_text != "":
                        document.add_paragraph(con_text)
                    for subsection in section["subsections"]:
                        document.add_heading(f"{subsection["subsection_title"]}", level=4)
                        scon_text = ""
                        for scon in subsection["subsection_content"]:
                            scon_text = scon_text + f"{scon["sentence"]} "
                        document.add_paragraph(scon_text)
                
                document.add_heading("References", level=3)
                
                for reference in essay["references"]:
                    document.add_paragraph(reference)
                    
                f = BytesIO()
                document.save(f)
                f.tell()
                f.seek(0)
                
                response = make_response(f.read())
                response.headers['Content-Type'] = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                response.headers['Content-Disposition'] = 'attachment; filename={}.docx'.format(document_obj.title)
                return response
        else:
            return Response(
                response=json.dumps({
                    "status": False,
                    "message": "Document not found"
                }),
                status=404,
                mimetype="application/json"
            )
            
    except Exception as e:
        return Response(
                response=json.dumps({
                    "status": False,
                    "message": "Error Occured",
                    "error": str(e)
                }),
                status=500,
                mimetype="application/json"
            )
        
@documents.route('/<int:document_id>/generate', methods=['GET'])
@authentication_required
def generate_document(document_id):
    try:
        document_obj = Document.query.get(document_id)
        
            
    except Exception as e:
        return Response(
                response=json.dumps({
                    "status": False,
                    "message": "Error Occured",
                    "error": str(e)
                }),
                status=500,
                mimetype="application/json"
            )